"""
build_paper_v2.py
=================
Build the IEEE Software feature article (.docx) from results_v2/.

Addresses every blocker from the v1 review:
  1. Section dispatcher uses exact prefix match (no more duplicate "Why" headings).
  2. Includes a 150-word formal Abstract.
  3. Includes the three Actionable Insights for Practitioners.
  4. Embeds only Figures 1, 2, 6 + the main results table; Figs 3-5 stay
     in results_v2/figures/ as supplemental material.
  5. Adds a legal-conflict sidebar (GDPR/AML, UAE erasure vs retention).
  6. Renders Markdown inline formatting (**bold**, *italic*, `code`) as
     real Word runs instead of leaving raw asterisks in the document.
  7. References are a proper numbered list with full IEEE-style details
     and in-text [N] citation markers.
  8. Softens over-strong claims per reviewer feedback.
  9. Adds Data and Code Availability + Funding + Conflict-of-interest notes.
"""

from __future__ import annotations

import csv
import os
import re
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULTS = os.path.join(HERE, "results_v2")
FIGURES = os.path.join(RESULTS, "figures")
PAPER_DIR = os.path.join(HERE, "paper")


# ---------------------------------------------------------------------------
# CSV helpers
# ---------------------------------------------------------------------------


def _read_summary() -> List[Dict[str, str]]:
    with open(os.path.join(RESULTS, "summary.csv"), newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def _read_conflicts() -> List[Dict[str, str]]:
    p = os.path.join(RESULTS, "conflicts.csv")
    if not os.path.exists(p):
        return []
    with open(p, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def _macro(rows: List[Dict[str, str]]) -> Dict[str, float]:
    """Macro-average over datasets with non-zero ground truth."""
    keys = ["sparql_P", "sparql_R", "sparql_F1",
            "sim_P", "sim_R", "sim_F1",
            "ent_P", "ent_R", "ent_F1",
            "tfidf_F1", "keyword_F1"]
    out = {k: 0.0 for k in keys}
    eligible = [r for r in rows if int(r.get("gt_conflicts", 0) or 0) > 0]
    n = max(len(eligible), 1)
    for r in eligible:
        for k in keys:
            out[k] += float(r.get(k, 0) or 0)
    for k in keys:
        out[k] /= n
    out["__n_datasets__"] = float(len(eligible))
    return out


# ---------------------------------------------------------------------------
# Inline-Markdown parser: **bold**, *italic*, `code`
# ---------------------------------------------------------------------------


# Order matters: try ** before * because ** would otherwise be eaten by *
_INLINE_PATTERN = re.compile(
    r"(\*\*([^*]+)\*\*)"  # **bold**
    r"|(`([^`]+)`)"        # `code`
    r"|(\*([^*]+)\*)"      # *italic*
)


def _parse_inline(text: str) -> List[Tuple[str, Dict[str, bool]]]:
    """
    Split a single line of body text into a list of (substring, flags) where
    flags is one of:
        {"bold": True}, {"italic": True}, {"code": True}, or {} (plain).
    """
    runs: List[Tuple[str, Dict[str, bool]]] = []
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], {}))
        if m.group(2) is not None:        # **bold**
            runs.append((m.group(2), {"bold": True}))
        elif m.group(4) is not None:      # `code`
            runs.append((m.group(4), {"code": True}))
        elif m.group(6) is not None:      # *italic*
            runs.append((m.group(6), {"italic": True}))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], {}))
    return runs


def _add_runs(paragraph, text: str, base_size: int = 11):
    """Add runs to a paragraph, parsing inline markdown."""
    for chunk, flags in _parse_inline(text):
        run = paragraph.add_run(chunk)
        run.font.size = Pt(base_size)
        if flags.get("bold"):
            run.bold = True
        if flags.get("italic"):
            run.italic = True
        if flags.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)


# ---------------------------------------------------------------------------
# Word formatting primitives
# ---------------------------------------------------------------------------


COLOR_TITLE = RGBColor(0x1f, 0x3b, 0x73)
COLOR_INK = RGBColor(0x22, 0x22, 0x22)
COLOR_DIM = RGBColor(0x55, 0x55, 0x55)


def _add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = COLOR_TITLE


def _add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_TITLE


def _add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = COLOR_TITLE


def _add_para(doc: Document, text: str, size: int = 11):
    p = doc.add_paragraph()
    _add_runs(p, text, base_size=size)


def _add_italic(doc: Document, text: str, size: int = 11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(size)
    r.font.color.rgb = COLOR_DIM


def _add_bullets(doc: Document, items: List[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _add_runs(p, item, base_size=11)


def _add_numbered(doc: Document, items: List[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        _add_runs(p, item, base_size=11)


def _add_sidebar(doc: Document, title: str, body_paragraphs: List[str]):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF4F4")
    cell._tc.get_or_add_tcPr().append(shd)

    p_title = cell.paragraphs[0]
    r = p_title.add_run("Sidebar: " + title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_TITLE
    for body in body_paragraphs:
        p = cell.add_paragraph()
        _add_runs(p, body, base_size=10)
    cell.add_paragraph()


def _add_figure(doc: Document, filename: str, caption: str, width_inches: float = 6.2):
    path = os.path.join(FIGURES, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_DIM


def _add_results_table(doc: Document, rows: List[Dict[str, str]]):
    table = doc.add_table(rows=1 + len(rows), cols=8)
    table.style = "Light Grid Accent 1"
    hdr = ["Dataset", "Rules", "Pairs", "GT conflicts (SMT oracle)",
           "F1 (full pipeline)", "F1 (TF-IDF baseline)", "F1 (keyword baseline)",
           "End-to-end time (s)"]
    for i, h in enumerate(hdr):
        c = table.rows[0].cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for ri, r in enumerate(rows, start=1):
        cells = [
            r["dataset"], r["rules"], r["total_pairs"], r["gt_conflicts"],
            r["sparql_F1"], r["tfidf_F1"], r["keyword_F1"], r["total_time_s"],
        ]
        for i, val in enumerate(cells):
            table.rows[ri].cells[i].text = str(val)
            for run in table.rows[ri].cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Manuscript content (markdown source for sanity / tracked-changes review)
# ---------------------------------------------------------------------------


def render_markdown(rows, conflicts, macro) -> str:
    """Plain markdown source mirroring the .docx body."""
    by_name = {r["dataset"]: r for r in rows}
    cont_a = by_name.get("Continue-A", {})
    synth = by_name.get("Synthetic360", {})
    geys = by_name.get("GEYSERS", {})
    kmkt = by_name.get("KMarket", {})

    total_smt_calls = sum(int(r.get("smt_total", 0) or 0) for r in rows)
    total_unknown = sum(int(r.get("smt_unknown", 0) or 0) for r in rows)
    cov = (total_smt_calls - total_unknown) / total_smt_calls * 100 if total_smt_calls else 100.0
    total_time = sum(float(r.get("total_time_s", 0) or 0) for r in rows)

    md = f"""# Don't Trust Similarity, Verify It: Solver-Backed Witnesses for Multi-Framework Policy Conflicts

Rabea Al Haj Eid, Mu'awya Al-Dala'ien (Princess Sumaya University for Technology, Amman, Jordan); Mohammad Al Haj Eid (Higher Colleges of Technology, Fujairah, UAE).

## Abstract

Organisations subject to overlapping security and privacy frameworks (NIST CSF, ISO/IEC 27001, PSD2, DORA, GDPR, and UAE AML/CFT obligations) routinely run rules that contradict each other at run time, yet auditors detect these reachable contradictions only after incidents. We present a four-stage pipeline that normalises heterogeneous XACML rules into a Common Policy Model, screens candidate pairs with sentence-embedding similarity, filters them with entity overlap and a SPARQL structural check, and verifies the survivors with the Z3 SMT solver. Every reported conflict is accompanied by a concrete witness request that an auditor can replay against the live policy decision point. On four public XACML datasets, the funnel narrows tens of thousands of pairs to seconds of solver work, with 100 percent SMT coverage (no unknowns, no timeouts) across 6,988 calls. The contribution is the funnel, not any single stage.

## Actionable Insights for Practitioners

- Treat semantic similarity as triage, not proof — pair it with structural filters and a solver before any conflict claim.
- Require every reported conflict to ship with a replayable witness request the auditor can dispatch against the live policy decision point.
- Put solver-backed policy checks into CI before policy-as-code changes reach production; a regression from `unsat` to `sat` should fail the build the same way a unit test does.

## Why this matters

Compliance teams rarely answer to one framework. A bank in Europe lives under PSD2 [3] and DORA [4]. A health provider in the US may use the NIST Cybersecurity Framework [1] alongside HIPAA-driven controls. A multinational adds ISO/IEC 27001 [2] on top of either, then has a board-driven internal control catalogue on top of that. Each framework normally has its own access-control rules, often written in slightly different vocabularies, and those rules end up enforced together by the same authentication, authorisation, and policy-decision services.

That overlap is where reachable contradictions get born. One rule from the PSD2 set says a customer may read their own account data when authenticated. An internal control says reads of account data are denied outside business hours. An auditor working from a control-mapping spreadsheet does not catch that conflict, because the spreadsheet only tells you that both controls are about "access to account data." A semantic-similarity model trained on policy text will tell you the two rules look alike, which is true but not actually useful — looking alike is not the same as colliding at run time. What an organisation actually needs is proof that *some real request* triggers both rules and gets contradictory answers, plus the request itself in a form that an engineer can replay against the live system.

We built a pipeline that produces exactly that: a concrete request, generated automatically, that simultaneously triggers two policy rules with opposite effects. The headline shape is deliberately conservative — cheap filters first, decisive proof last. There is no machine-learning classifier in the conflict-detection loop. Every conflict the pipeline reports is accompanied by a witness request that an auditor or compliance engineer can run through the real Policy Decision Point (PDP) and watch the contradiction happen.

This article walks through what we built, what surprised us when we measured it (including a ground-truth-counting bug that we caught and want to share publicly), and the practitioner takeaways for anyone shipping policy-as-code today.

## Why now

Two things changed in the last few years that made this work tractable. First, regulated-services frameworks moved from PDF-and-spreadsheet artefacts toward machine-readable forms — XACML 3.0 [7], OSCAL, and the recent practice of treating internal control catalogues as version-controlled YAML or JSON. That gives us a structured input format, which is the precondition for any kind of automated reasoning over policy. Second, SMT solvers got fast enough on the kinds of formulas access-control rules produce that you can call them in a tight loop on a laptop. Z3 [6] returning a sat/unsat answer in single-digit milliseconds for a typical XACML rule pair means a verification stage is no longer a quarterly audit ritual — it is something the build can do on every pull request.

What did not change is the gap between control mappings and run-time enforcement. Auditors still spend most of their time correlating control-catalogue entries from one framework to another, and most of those correlations are valid at the *intent* level but say nothing about whether the actual XACML or Rego [8] rules implementing those intents will collide at run time. The pipeline we describe is meant to live inside that gap: it takes the rules as they are actually written, normalises them, and tells you, with a witness or with a proof, whether two of them collide.

A practitioner reading this should care because the alternative is one of two failure modes that we have both seen in the wild. One is shipping a similarity-based "AI policy linter" that produces confident false positives and trains operators to ignore it. The other is leaving the cross-framework reasoning to spreadsheet triage, which works until a regulator asks for a defensible answer about a specific reachable contradiction. Solver-backed witnesses dodge both failure modes by giving you the artefact regulators ask for: a concrete request that demonstrates the conflict.

## What we built

The pipeline has three filtering stages followed by one solver stage. Figure 1 shows the funnel.

The Common Policy Model (CPM) normalises rules from XACML files [7], NIST/ISO control catalogues [1], [2], PSD2/DORA mappings [3], [4], and internal documentation into a tuple of subject attributes, action verb, resource type and identifiers, environment conditions, effect (Permit or Deny), framework origin, and priority. Semantic screening uses a sentence-transformer [15] to score every pair of rules for textual similarity; pairs below the chosen threshold are dropped. The entity-overlap stage uses Jaccard similarity on extracted subject roles and resource identifiers, and the SPARQL stage queries an RDF view [14] of the policy graph to confirm that the two rules can plausibly reach the same scope. Finally, every pair that survives is encoded as a satisfiability problem and handed to Z3 [6]. If Z3 returns `sat`, the model is a concrete request that triggers both rules; if `unsat`, the pair is provably non-conflicting.

The point is not that any single stage is novel. Sentence transformers, RDF/SPARQL, and SMT-backed XACML analysis [9], [10], [11], [12] have all been done separately. The contribution is the funnel: the filters cut solver work by orders of magnitude, and the solver is the only stage we trust for ground truth. Detection without verification gives you opinions; verification without filtering gives you a stalled pipeline.

### How it works in practice

Each rule, after CPM normalisation, carries a set of attribute constraints inherited down the XACML PolicySet -> Policy -> Rule hierarchy. The hierarchy matters: a rule's own `<Target>` element is often empty, and the actual scope is supplied by the parent Policy and PolicySet. We walk the tree, accumulate constraints, and respect XACML's `AnyOf`/`AllOf` semantics — `AnyOf` is a disjunction, `AllOf` is a conjunction, and a rule's full applicability formula is the AND of its inherited `AnyOf` groups.

Two rules conflict iff there exists an assignment to the attribute variables in the request `q` such that the applicability formulas of both rules are simultaneously satisfied **and** the rules disagree on effect. The SMT encoder turns each rule's applicability into a Z3 formula, with a single fresh variable per attribute (so equality and comparison constraints on the same attribute interact properly). Same-effect pairs short-circuit instantly. Cross-effect pairs are handed to Z3, which either returns a model or a proof of unsatisfiability.

A few engineering details matter in practice. We share the Z3 variable environment between the two rules in a conflict check, so equality and comparison constraints on the same attribute interact correctly inside the solver — earlier versions of our encoder created independent variables for `string-equal` and `string-less-than-or-equal` constraints, which silently dropped the interaction and let pairs slip through as false negatives. We pre-compute per-rule scope satisfiability once before running the cross-product, so a rule whose own constraints are unsatisfiable (dead-code under the current XACML hierarchy) short-circuits every pair it participates in. And we deduplicate candidate pairs before the SMT step, since the upstream filters can emit the same pair more than once.

### How the funnel relates to prior tools

Researchers have been chasing pieces of this problem for over a decade. Margrave [11] used MTBDDs to analyse firewall configurations; XEngine [9] was a fast XACML evaluation engine; ACPT [10] modelled and verified XACML policies; Turkmen et al. [12] put XACML into SMT directly; Caserio et al. [13] formally validated XACML policies. Each of those tools is good at its layer of the problem, but each treats a single policy set as the input. The cross-framework problem is where the existing tooling falls short, because nobody normalised the inputs first. The CPM is the layer where we make the inputs comparable; the funnel is what makes verification affordable across the combined input.

The honest comparison is not "we beat Margrave's runtime." It is: the prior tools were right about the value of formal analysis, but limited to one policy set at a time; our contribution is to put that formal analysis behind a similarity-and-graph funnel so it can run across heterogeneous standards in a build pipeline.

## What we measured

We evaluated the pipeline on four public XACML datasets:

- **GEYSERS** ({geys.get('rules','15')} rules, EU cloud infrastructure project),
- **KMarket** ({kmkt.get('rules','5')} rules, e-commerce marketplace),
- **Continue-A** ({cont_a.get('rules','298')} rules, healthcare access control),
- **Synthetic360** ({synth.get('rules','360')} rules, synthetic benchmark).

Ground truth comes from running Z3 on every cross-effect rule pair. Same-effect pairs cannot collide on effect, so we skip them. For Continue-A this is {cont_a.get('total_pairs','44253')} total pairs, of which {cont_a.get('gt_sat','482')} are real conflicts under the SMT oracle. For Synthetic360, every rule's accumulated scope under the strict XACML hierarchy turned out to be unsatisfiable on its own — the synthetic generator stacks so many `AnyOf` groups that no concrete request can satisfy them all. That dataset has zero reachable conflicts in the strict sense, which is itself an interesting finding: synthetic stress-test data does not always behave like real policy.

### Headline numbers

The full pipeline reaches macro Precision = {macro['sparql_P']:.2f}, Recall = {macro['sparql_R']:.2f}, F1 = {macro['sparql_F1']:.2f} against the SMT oracle across the datasets that have non-zero ground truth (KMarket and Continue-A). On Continue-A specifically, the funnel narrows from 44,253 raw pairs down to 70 candidates, and SMT confirms 55 of those 70 as real conflicts (a per-call precision of 0.79 at the verification step). The solver returned a definite answer on every pair we handed it — no `unknown` outcomes, no timeouts — across {total_smt_calls} calls in total. End-to-end runtime for all four datasets together was {total_time:.0f} seconds on a Windows 11 laptop with no GPU. Z3 [6] was reliable for the XACML-shaped rules in our evaluation; we do not generalise that claim to richer schemas with non-linear arithmetic.

Two things to read here. First, the detector is moderate. Similarity-only retrieves a lot of false positives because rule text reuses the same vocabulary; entity overlap and SPARQL drop those, but recall pays a price. Second, the verifier is decisive. There is no "we think this is a conflict, confidence 0.83" output anywhere in the pipeline. Every reported conflict has a witness request attached and every cleared pair has a `unsat` proof. That is the contribution: not a strong detector, but a detector deliberately calibrated to feed a strong verifier.

The threshold sweep (Figure 2) shows that F1 is fairly flat in a 0.5-0.7 band, which is a good sign for transferability — you are not balanced on a knife-edge. The main results are in Table 1; supplemental figures (filter cascade, SMT outcome distribution per dataset, runtime breakdown) are in `results_v2/figures/` in the public repository [16].

## What we learned

**1. Similarity is not enough.** Two policy rules can have ~0.9 cosine similarity and still target completely different scopes. The reverse also happens: two rules with ~0.4 similarity can be hard contradictions because the wording is different but the underlying semantics collapse. The semantic stage is best understood as a recall-preserving funnel, not as a classifier.

**2. The SPARQL filter is cheap and load-bearing.** It cut solver candidates by roughly an order of magnitude on Continue-A. Without it the SMT stage still works, but with it the pipeline runs in seconds rather than tens of seconds on a laptop.

**3. Z3 on XACML-shaped rules is fast and decisive in our evaluation.** Median solve time per pair was a few milliseconds. We observed zero `unknown` outcomes and zero timeouts under a five-second per-call budget on the four datasets we tested. This will not scale to richer attribute-based access-control schemas with non-linear arithmetic or recursion over policy sets, but for the equality-and-comparison scope of typical XACML rules, the solver was a tool we could plan budget around.

**4. Ground truth by arithmetic is a trap.** Our first labeller counted "every Permit-Deny pair where the rules share resource type" as a conflict. On Continue-A, where every rule shares the same resource type, that collapsed to "every Permit-Deny pair." The number it produced — 14,280 — was exactly Permits (238) times Denies (60). On Synthetic360 the same pattern produced 32,399, exactly 179 times 181. We caught it because both numbers factored too cleanly. The fix was to use the SMT result as the source of truth. The lesson generalises: any "ground truth" that is an arithmetic identity over your input is suspect.

**5. Synthetic stress-test data does not always behave like real policy.** The Synthetic360 dataset stacks so many `AnyOf` constraints that under strict XACML semantics every rule's scope is unsatisfiable on its own. The verifier correctly identifies this; under the original wildcard-leaky encoder, every cross-effect pair looked like a conflict.

**6. Witness-as-request makes audit defensible.** When a reviewer asks how we know a conflict is real, the witness gives reviewers concrete evidence to inspect — a request they can dispatch against the live PDP and observe two enforcement answers come back. We have found witness rendering to be the single most important interface decision in the pipeline.

### Threats to validity

Two we want to flag up front. First, our datasets are public XACML benchmarks; they are not the kind of cross-framework policy mix that a real bank or healthcare organisation operates. The conflict-detection numbers we report are honest measurements against a strict XACML oracle on these datasets, but they should not be over-generalised. Second, the SMT encoder currently ignores XACML obligations and advice. A pair that the solver labels `unsat` could still produce contradictory obligations at run time even when the effect is consistent, and we do not yet model that. Both threats are limitations of scope rather than of the approach.

We also want to be transparent about a non-finding: the detection F1 we report is moderate (around 0.30 macro across the two datasets with non-zero ground truth), and we are not claiming the detector is the contribution. Earlier versions of this work reported much higher F1 numbers against a permissive ground-truth definition that turned out to be the bug we describe in the GT-counting sidebar. With the corrected oracle, the honest detection numbers are lower — but the verification step is what compliance teams need, and the verification step is decisive.

## What's next

The current encoding handles equality and basic comparison operators on string and numeric attributes. Three pragmatic next steps:

- **Regex and obligation reasoning.** XACML's regex predicates and obligation/advice machinery are not yet in the encoder. Both are pragmatic additions; neither requires a fundamentally different solver back-end.
- **CI gate.** A policy-as-code change that introduces an `unsat` -> `sat` regression on a previously cleared pair should fail the build the same way a unit test does. We are working on a thin GitHub Action that runs the pipeline on changed rules and posts the witness as a PR comment.
- **Live PDP integration.** The witness request is currently a text artefact; the obvious next step is to dispatch it against the live PDP automatically and show the auditor the two enforcement answers side by side.

For practitioners shipping policy-as-code today, the takeaway is simpler. Do not trust a similarity model to tell you whether two rules contradict. Use cheap filters to cut the search space, then let a solver give you a witness or a proof of non-conflict. The solver is fast enough on XACML-shaped rules for it to be a routine part of the build, not a quarterly audit ritual. And whenever a "conflict count" looks like a clean factorisation of two rule subtotals, look again — you may be measuring an arithmetic identity instead of a fact about your policy.

### How to adopt this in your own pipeline

Start with normalisation: pick or invent a Common Policy Model that covers every framework you care about, and write the importers — XACML, Rego, OSCAL, internal YAML — so they all land in the same shape. Going broad and shallow on the CPM beats going narrow and deep — the marginal correctness from a richer CPM is usually less than the marginal coverage you get from supporting a third or fourth framework.

Next, plug in a sentence-transformer for similarity screening. We used `all-MiniLM-L6-v2` because it is small, fast on CPU, and good enough for the triage role. Tune the threshold on a validation slice of your own data; the 0.65 we report is a reasonable starting point.

Then add the structural filters. Entity overlap is the cheap one and gives you most of the precision improvement; a SPARQL or graph-query layer adds another order-of-magnitude reduction in solver candidates if you can afford the engineering cost of an RDF view.

Finally, run Z3 on the survivors with a fixed per-call budget. Render every `sat` model as a request the auditor can replay. Persist the witnesses; they are the only artefact you need to defend a conflict claim to a regulator. Treat unsat-scope rules as dead code and surface them to the rule author. And if you are tempted to write your own "ground-truth" labeller for evaluation, do not — the SMT result is the ground truth.

## Data and code availability

All code, datasets (XACML files), v2 pipeline, figure-generation scripts, and the IEEE Software manuscript source are available at the public repository [16]. The four datasets used in this article are public XACML benchmarks distributed with prior work in the field. The corrected v2 pipeline reproduces every number reported in this article in approximately 100 seconds on a Windows 11 laptop without a GPU.

## Funding

This work received no external funding.

## Conflict of interest

The authors declare no conflict of interest.

## The Authors

**Rabea Al Haj Eid** is a research student at Princess Sumaya University for Technology in Amman, Jordan. He works on policy reasoning and applied formal methods for governance, risk, and compliance pipelines. Contact: rab20248091@std.psut.edu.jo.

**Mu'awya Al-Dala'ien** is a faculty member at Princess Sumaya University for Technology in Amman, Jordan. His research focuses on cybersecurity governance, regulated digital services, and the engineering of compliance automation. Contact: m.aldalaien@psut.edu.jo.

**Mohammad Al Haj Eid** is a faculty member at the Higher Colleges of Technology in Fujairah, UAE. His work spans secure software engineering, access-control automation, and applied AI for cybersecurity. Contact: Malhajeid@hct.ac.ae.
"""
    return md


# ---------------------------------------------------------------------------
# Reference list (numbered, full IEEE-style details)
# ---------------------------------------------------------------------------


REFERENCES: List[str] = [
    "National Institute of Standards and Technology, \"The NIST Cybersecurity Framework (CSF) 2.0,\" NIST Cybersecurity White Paper CSWP 29, Feb. 2024.",
    "International Organization for Standardization, \"ISO/IEC 27001:2022 - Information security, cybersecurity and privacy protection - Information security management systems - Requirements,\" ISO/IEC, Geneva, 2022.",
    "European Union, \"Directive (EU) 2015/2366 of the European Parliament and of the Council of 25 November 2015 on payment services in the internal market (PSD2),\" Official Journal of the European Union, L 337, pp. 35-127, 2015.",
    "European Union, \"Regulation (EU) 2022/2554 on digital operational resilience for the financial sector (DORA),\" Official Journal of the European Union, L 333, pp. 1-79, 2022.",
    "European Union, \"Regulation (EU) 2016/679 on the protection of natural persons with regard to the processing of personal data and on the free movement of such data (GDPR),\" Official Journal of the European Union, L 119, pp. 1-88, 2016.",
    "L. de Moura and N. Bjorner, \"Z3: An Efficient SMT Solver,\" in Tools and Algorithms for the Construction and Analysis of Systems (TACAS), LNCS 4963, pp. 337-340, Springer, 2008.",
    "OASIS, \"eXtensible Access Control Markup Language (XACML) Version 3.0, Plus Errata 01,\" OASIS Standard Incorporating Approved Errata, Jul. 2017.",
    "Open Policy Agent project, \"Rego policy language,\" https://www.openpolicyagent.org, accessed Apr. 2026.",
    "A. X. Liu, F. Chen, J. Hwang, and T. Xie, \"XEngine: A fast and scalable XACML policy evaluation engine,\" in Proc. ACM SIGMETRICS Int. Conf. Measurement and Modeling of Computer Systems, pp. 265-276, 2008.",
    "J. Hwang, T. Xie, V. Hu, and M. Altunay, \"ACPT: A tool for modeling and verifying access control policies,\" in Proc. IEEE Int. Symp. Policies for Distributed Systems and Networks (POLICY), pp. 40-43, 2010.",
    "T. Nelson, C. Barratt, D. J. Dougherty, K. Fisler, and S. Krishnamurthi, \"The Margrave Tool for Firewall Analysis,\" in Proc. USENIX LISA, pp. 65-76, 2010.",
    "F. Turkmen, J. den Hartog, S. Ranise, and N. Zannone, \"Formal analysis of XACML policies using SMT,\" Computers & Security, vol. 66, pp. 185-203, May 2017.",
    "C. Caserio, F. Lonetti, and E. Marchetti, \"Formal validation of XACML policies,\" Software Quality Journal, vol. 30, pp. 877-907, 2022.",
    "World Wide Web Consortium (W3C), \"RDF 1.1 Concepts and Abstract Syntax\" and \"SPARQL 1.1 Query Language,\" W3C Recommendations, 2014.",
    "N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence embeddings using Siamese BERT-networks,\" in Proc. EMNLP-IJCNLP, pp. 3982-3992, 2019.",
    "Hellscream999, \"Policy Conflict Detection (corrected v2 pipeline),\" GitHub repository, https://github.com/Hellscream999/policy-conflict-detection, 2026.",
    "United Arab Emirates Federal Decree-Law No. 45 of 2021, \"On the Protection of Personal Data,\" UAE Official Gazette, 2021.",
    "Central Bank of the United Arab Emirates, \"AML/CFT Compliance Standards,\" Rulebook 16, including 16.29.1 \"Record Retention,\" various editions.",
]


# ---------------------------------------------------------------------------
# Section dispatcher (exact-match keys, no accidental prefix collisions)
# ---------------------------------------------------------------------------


def _process_blocks(doc: Document, md_text: str, rows: List[Dict[str, str]],
                    conflicts: List[Dict[str, str]]) -> None:
    """
    Walk the markdown blocks and emit Word output. Each section's heading
    paragraph is matched on its exact text (not a prefix), so we never emit
    the same heading or figure twice.
    """
    inserted_fig1 = False
    inserted_legal_sidebar = False
    inserted_what_we_mean_sidebar = False
    inserted_table = False
    inserted_fig2 = False
    inserted_gt_sidebar = False
    inserted_witness_sidebar = False
    inserted_fig6 = False

    blocks = [b.strip() for b in md_text.split("\n\n") if b.strip()]

    for block in blocks:
        # ----- title (handled outside) -----
        if block.startswith("# "):
            continue

        # ----- author byline -----
        if block.startswith("Rabea Al Haj Eid"):
            _add_italic(doc, block)
            continue

        # ----- section dispatcher: exact heading text match -----
        if block == "## Abstract":
            _add_h2(doc, "Abstract")
            continue
        if block == "## Actionable Insights for Practitioners":
            _add_h2(doc, "Actionable Insights for Practitioners")
            continue
        if block == "## Why this matters":
            _add_h2(doc, "Why this matters")
            if not inserted_fig1:
                _add_figure(doc, "fig1_architecture.png",
                            "Figure 1. Funnel architecture: cheap filters first, decisive proof last. "
                            "Heterogeneous policies are normalised into the Common Policy Model, then "
                            "filtered by similarity, entity overlap, and SPARQL before reaching the SMT "
                            "verifier. Z3 returns either a concrete witness request or an unsat proof.")
                inserted_fig1 = True
            continue
        if block == "## Why now":
            # Insert the "What we mean by 'conflict'" sidebar between Why this
            # matters and Why now -- the conceptual definition belongs early.
            if not inserted_what_we_mean_sidebar:
                _add_sidebar(doc,
                    "What we mean by 'conflict'",
                    [
                        "Two policy rules conflict when there exists a single concrete request that "
                        "triggers both rules and the rules disagree on effect — one says Permit, the "
                        "other says Deny. The existence of that request is the test, not textual "
                        "similarity and not control-mapping overlap.",
                        "We classify reachable conflicts into four operational categories. **Exact "
                        "contradiction**: same scope, opposite effect. **Attribute conflict**: different "
                        "subject conditions trigger inconsistent enforcement. **Scope conflict**: "
                        "overlapping but mismatched conditions, where the overlap region is non-empty "
                        "but the rules' authors did not realise it. **Priority conflict**: incompatible "
                        "obligations or defaults that require an explicit precedence rule to resolve.",
                    ])
                inserted_what_we_mean_sidebar = True
            # And the legal-conflict sidebar belongs here too -- right after the
            # motivating section, anchored to real regulations.
            if not inserted_legal_sidebar:
                _add_sidebar(doc,
                    "A legal conflict becomes an engineering conflict",
                    [
                        "Cross-standard contradictions are not academic. The GDPR right to erasure "
                        "(Article 17 [5]) requires a controller to delete personal data on request. "
                        "EU AML and counter-terrorist-financing rules require regulated entities to "
                        "retain transaction records for a fixed minimum period. Both apply to the same "
                        "natural-person record at the same time. The auditor sees two compatible "
                        "intents on a control mapping; the policy decision point sees one Permit-erase "
                        "and one Deny-erase against the same subject and resource.",
                        "The same shape recurs in the United Arab Emirates. Federal Decree-Law No. 45 "
                        "of 2021 [17] gives data subjects an erasure right. Central Bank AML/CFT "
                        "Rulebook clauses such as 16.29.1 [18] require record retention for fixed "
                        "periods. Implemented as XACML or Rego rules, they collide on a single "
                        "(subject, action=delete, resource=customer-record) request.",
                        "The pipeline's contribution to these cases is not a legal opinion — it is a "
                        "concrete request that demonstrates the collision. The auditor still has to "
                        "decide which obligation wins, but the engineering question (does this collide "
                        "at run time?) gets a yes/no answer with a replayable witness, instead of a "
                        "spreadsheet correlation.",
                    ])
                inserted_legal_sidebar = True
            _add_h2(doc, "Why now")
            continue
        if block == "## What we built":
            _add_h2(doc, "What we built")
            continue
        if block == "### How it works in practice":
            _add_h3(doc, "How it works in practice")
            continue
        if block == "### How the funnel relates to prior tools":
            _add_h3(doc, "How the funnel relates to prior tools")
            continue
        if block == "## What we measured":
            _add_h2(doc, "What we measured")
            continue
        if block == "### Headline numbers":
            _add_h3(doc, "Headline numbers")
            # Insert the main results table right after the headline numbers.
            if not inserted_table:
                _add_results_table(doc, rows)
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run("Table 1. Per-dataset results from the v2 pipeline. F1 is "
                                  "measured against the SMT-oracle ground truth. End-to-end time "
                                  "covers all four stages including model load.")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_DIM
                inserted_table = True
            # And Figure 2 (threshold sweep) here, since this is where the
            # "where to set the threshold" discussion belongs.
            if not inserted_fig2:
                _add_figure(doc, "fig2_threshold_sweep.png",
                            "Figure 2. F1 versus the similarity threshold and the precision-recall "
                            "trace per dataset. F1 is stable in the 0.5-0.7 band, indicating that "
                            "the funnel is not balanced on a knife-edge.")
                inserted_fig2 = True
            continue
        if block == "## What we learned":
            _add_h2(doc, "What we learned")
            continue
        if block == "### Threats to validity":
            # Drop the GT-counting sidebar between What we learned and Threats
            # to validity. It is the most useful debrief in the article.
            if not inserted_gt_sidebar:
                _add_sidebar(doc,
                    "The ground-truth bug we caught (and why we're telling you)",
                    [
                        "Our first ground-truth labeller marked a pair a conflict whenever the two "
                        "rules had opposite effects and shared a resource type or action. The intent "
                        "was to catch every obvious collision; the implementation was too forgiving. "
                        "On Continue-A, where every rule happens to share the same resource type, "
                        "the rule collapsed to: every Permit-Deny pair is a conflict. The reported "
                        "count was 14,280 — exactly Permits (238) times Denies (60). On Synthetic360, "
                        "32,399 — exactly 179 times 181.",
                        "We caught it because two independent datasets produced two clean "
                        "factorisations on the nose. Real conflict counts have noise; arithmetic "
                        "identities do not. The fix was to make the Z3 result the source of truth — "
                        "the conflict-existence test was already defined as 'a single request "
                        "triggers both rules with opposite effects', which is precisely what an SMT "
                        "instance answers. Once we ran the solver as the labeller, the Continue-A "
                        "count dropped to 482 (1.09 percent of pairs) and Synthetic360 dropped to "
                        "zero (because every Synthetic360 rule's accumulated scope is unsatisfiable "
                        "under strict XACML semantics).",
                        "The general lesson: any conflict count that is an exact product of two "
                        "rule-count subtotals deserves a second look. Arithmetic identities can "
                        "masquerade as data.",
                    ])
                inserted_gt_sidebar = True
            _add_h3(doc, "Threats to validity")
            continue
        if block == "## What's next":
            _add_h2(doc, "What's next")
            continue
        if block == "### How to adopt this in your own pipeline":
            _add_h3(doc, "How to adopt this in your own pipeline")
            continue
        if block == "## Data and code availability":
            # Witness sidebar + Figure 6 belong before the closing matter.
            if not inserted_witness_sidebar:
                _add_sidebar(doc,
                    "Reading a Z3 witness",
                    [
                        "When the solver returns sat, it produces a model — concrete values for "
                        "every Z3 variable in the conflict formula. We bind those values back to the "
                        "CPM attributes they came from (subject, role, action, resource, environment "
                        "conditions) and present the result as a request. A real example from "
                        "Continue-A: subject=pc-member, action=read, resource=paper-review_rc. That "
                        "single request triggers two policy rules — one says Permit, the other Deny — "
                        "and the request is concrete enough to dispatch.",
                        "An auditor can run that request through the live PDP and see two enforcement "
                        "answers come back. That is what 'reachable' means in operational terms. The "
                        "witness gives reviewers concrete evidence to inspect, instead of a "
                        "confidence score they can argue with.",
                    ])
                inserted_witness_sidebar = True
            if not inserted_fig6:
                _add_figure(doc, "fig6_witness.png",
                            "Figure 3. A concrete Z3 witness from Continue-A rendered as a request, "
                            "with the two rule outcomes it triggers. Supplemental figures (filter "
                            "cascade, SMT outcome distribution, runtime breakdown) are in the public "
                            "repository [16].")
                inserted_fig6 = True
            _add_h2(doc, "Data and code availability")
            continue
        if block == "## Funding":
            _add_h2(doc, "Funding")
            continue
        if block == "## Conflict of interest":
            _add_h2(doc, "Conflict of interest")
            continue
        if block == "## The Authors":
            _add_h2(doc, "The Authors")
            continue

        # Bullet list block (Actionable Insights)
        if block.startswith("- "):
            items = [line[2:].strip() for line in block.split("\n") if line.startswith("- ")]
            _add_bullets(doc, items)
            continue

        # Plain paragraph - the inline parser handles **bold**, *italic*, `code`.
        # We collapse single newlines inside a paragraph into spaces.
        flat = " ".join(block.split())
        _add_para(doc, flat)


# ---------------------------------------------------------------------------
# Top-level build
# ---------------------------------------------------------------------------


def build_docx(out_dir: str):
    rows = _read_summary()
    conflicts = _read_conflicts()
    macro = _macro(rows)
    md_text = render_markdown(rows, conflicts, macro)

    doc = Document()

    # margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    # title
    _add_h1(doc,
            "Don't Trust Similarity, Verify It: Solver-Backed Witnesses "
            "for Multi-Framework Policy Conflicts")
    _add_italic(doc,
        "Rabea Al Haj Eid and Mu'awya Al-Dala'ien — Princess Sumaya University for "
        "Technology, Amman, Jordan")
    _add_italic(doc,
        "Mohammad Al Haj Eid — Higher Colleges of Technology, Fujairah, UAE")

    # body
    _process_blocks(doc, md_text, rows, conflicts)

    # references
    _add_h2(doc, "References")
    for i, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"[{i}] ")
        r.bold = True
        r.font.size = Pt(10)
        _add_runs(p, ref, base_size=10)

    # save
    os.makedirs(out_dir, exist_ok=True)
    docx_path = os.path.join(out_dir, "IEEE_Software_PolicyConflict_v2.docx")
    md_path = os.path.join(out_dir, "manuscript.md")
    doc.save(docx_path)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_text)

    body_words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"docx -> {docx_path}")
    print(f"md   -> {md_path}")
    print(f"approx body word count = {body_words}")


if __name__ == "__main__":
    build_docx(PAPER_DIR)
